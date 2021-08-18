from docx import Document
from docx.enum.style import WD_STYLE_TYPE # Define estilo para paragrafo
from docx.enum.table import WD_TABLE_ALIGNMENT # Define estilo para paragrafo
from docx.enum.text import WD_ALIGN_PARAGRAPH #Define alinhamento de paragráfo
from docx.shared import Pt #Define tamanho da fonte
from docx.shared import RGBColor #Define a cor da fonte
from docx.shared import Inches #Define tamanho
from typing import Counter
import matplotlib.pyplot as ps
import numpy as np
import os
from tkinter import filedialog
from tkinter import *
from datetime import date
import random
from PyQt5 import uic, QtWidgets, QtCore
from PyQt5.QtWidgets import QMessageBox, QWidget
import docx2pdf
#import win32com.client
from PyQt5.QtCore import Qt
import pyodbc
#import time


'''
numprocessos = 6
valor = {
    "Titulo":"Calculo de Escalonamento de Processos",
    "Informações": [
        {
            "TipoProcesso":"Processo selecionado: "+str(stre),
            "NumProcessos":str(numprocessos)+" Processos",
            "Image":"PipesTransparent.png",
            "Table":["A","b","c"],
        }
    ]
}
'''
def conexaobanco(nomeserver: str, nomebanco: str):
    #conexao = 'Driver={SQL Server Native Client 11.0};Server='+server+';Database='+database+';UID='+username+';PWD='+ password
    conexao = "Driver={SQL Server Native Client 11.0};Server="+nomeserver+";Database="+nomebanco+";Trusted_Connection=yes;"
    return pyodbc.connect(conexao)

def GerarRelatorio(processos, valuecalcu, gerartabela, tipoprocesso, tipocriterio) -> bool:
    ps.close()
    """[summary]

    Args:
        processos ([type]): [A variável Processos precisa ser carregada]
        valuecalcu ([type]): [O calculo de escalonamento]
        gerartabela ([type]): [GerarTabela(Calculo de escalonamento(com processo), False)]
        tipoprocesso ([type]): [FIFO / SJF / Por Prioridade]
        tipocriterio ([type]): [Nenhum / FIFO / SJF / Por Prioridade]
    """
    salvou = False
    valor = gerartabela
    # VARIAVEIS DE CONSTRUÇÃO DO GRÁFICO
    y = valor[0] # PROCESSOS DA TABELA Y
    c = []  # TIPO USADO (quantidade de processos)
    for x in y:
        c.append(1)
    xinicio = valor[1] #Inicio processo
    xfim = valor[2] # Fim processo
    color_mapper = np.vectorize(lambda x: {0: 'red', 1: 'blue'}.get(x))
    ps.hlines(y, xinicio, xfim, colors=color_mapper(c), lw=25)
    ps.ylabel('PROCESSOS')
    ps.xlabel('Tempo de Execução (ns)')
    ps.grid(True) #Deixar tabelado
    ps.savefig('image.png')

    docum = Document()

    docum.add_heading('Relatório: Cálculo de Escalonamento', 0)
    docum.add_heading('Tipo de escalonamento: {0}\nCritério de Desempate: {1}'.format(tipoprocesso, tipocriterio))
    styles = docum.styles
    # Style Paragraph
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    #Style Heading 2
    h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = "Calisto MT"
    h2.font.size = Pt(12)
    h2.font.bold = True #Define se vai ser negrito ou não
    #Style Heading 3
    h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = "Calisto MT"
    h3.font.size = Pt(10)
    h3.font.color.rgb = RGBColor(11, 0x0A, 0x6A) ##110A6A
    h3.font.bold = False #Define se vai ser negrito ou não

    tabela = docum.add_table(rows = len(processos)+1, cols = 4, style='Table Grid')
    Format = tabela.rows[0].cells
    Format[0].text = 'Processos'
    Format[1].text = 'Tempo de Execução'
    Format[2].text= 'Ordem de Chegada'
    Format[3].text= 'Prioridade'
    tabela.aligment = WD_TABLE_ALIGNMENT.CENTER
    for x in range(1, len(processos)+1):
        Format = tabela.rows[x].cells
        Format[0].text = str(processos[x-1][0])
        Format[1].text = str(processos[x-1][1])
        Format[2].text = str(processos[x-1][2])
        Format[3].text = str(processos[x-1][3])
    docum.add_paragraph('')
    docum.add_picture('image.png', width=Inches(4))
    val = 0
    for x in y:
        docum.add_paragraph('Tempo de Espera do Processo '+str(x)+': '+str(xinicio[val])+' NanosSegundos', style="H3")
        val+=1
    docum.add_paragraph('Tempo médio de espera: {0:.2f} ns'.format(valuecalcu[4]), style="H2")
    docum.add_paragraph('Tempo Total: {0:.2f} ns'.format(valuecalcu[2]), style="H2")
    Diretorio = Tk()
    try:
        Diretorio.withdraw()
        DiretorioSelecionado = filedialog.askdirectory()
        data = str(date.today())
        codigorandom = str(random.randrange(6000, 9999))
        if DiretorioSelecionado:
            dirfinaldocx = DiretorioSelecionado+'\RelatorioPIPE-'+data+'-'+codigorandom+'.docx'
            dirfinalpdf = DiretorioSelecionado+'\RelatorioPIPE-'+data+'-'+codigorandom+'.pdf'
            docum.save(dirfinaldocx)
            #docx2pdf.convert(r"{0}".format(dirfinaldocx), r"{0}".format(dirfinalpdf))
            ##################################################################################
            docx2pdf.convert(r"{0}".format(dirfinaldocx), r"{0}".format(dirfinalpdf))
            '''
            wdFormatPDF = 17
            in_file=dirfinaldocx
            out_file=dirfinalpdf
            word=win32com.client.DispatchEx("Pdf.Application")
            #word.Visible = True
            doc=word.Documents.Open(in_file) # Abre o arquivo docx
            doc.SaveAs(out_file, FileFormat=wdFormatPDF) # Faz a conversão em PDF
            doc.Close() # Fecha o arquivo .docx
            #word.Visible = False
            word.Quit() # Fechar aplicação
            '''
            #################################################################################
            os.remove(dirfinaldocx)
            salvou = True
    except Exception as error:
        salvou = False
        print(error)
    os.remove('image.png')
    ps.close()
    return salvou

def GerarRelatorioSQL(processos, valuecalcu, gerartabela, tipoprocesso, tipocriterio, processo = None) -> bool:
    ps.close()
    """[summary]

    Args:
        processos ([type]): [A variável Processos precisa ser carregada]
        valuecalcu ([type]): [O calculo de escalonamento]
        gerartabela ([type]): [GerarTabela(Calculo de escalonamento(com processo), False)]
        tipoprocesso ([type]): [FIFO / SJF / Por Prioridade]
        tipocriterio ([type]): [Nenhum / FIFO / SJF / Por Prioridade]
    """
    salvou = False
    valor = gerartabela
    codigo = criarcodigo()
    # VARIAVEIS DE CONSTRUÇÃO DO GRÁFICO
    y = valor[0] # PROCESSOS DA TABELA Y
    c = []  # TIPO USADO (quantidade de processos)
    for x in y:
        c.append(1)
    xinicio = valor[1] #Inicio processo
    xfim = valor[2] # Fim processo
    color_mapper = np.vectorize(lambda x: {0: 'red', 1: 'blue'}.get(x))
    ps.hlines(y, xinicio, xfim, colors=color_mapper(c), lw=25)
    ps.ylabel('PROCESSOS')
    ps.xlabel('Tempo de Execução (ns)')
    ps.grid(True) #Deixar tabelado
    ps.savefig('image.png')

    docum = Document()

    docum.add_heading('Relatório: Cálculo de Escalonamento', 0)
    docum.add_heading('Tipo de escalonamento: {0}\nCritério de Desempate: {1}'.format(tipoprocesso, tipocriterio))
    styles = docum.styles
    # Style Paragraph
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    #Style Heading 2
    h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = "Calisto MT"
    h2.font.size = Pt(12)
    h2.font.bold = True #Define se vai ser negrito ou não
    #Style Heading 3
    h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = "Calisto MT"
    h3.font.size = Pt(10)
    h3.font.color.rgb = RGBColor(11, 0x0A, 0x6A) ##110A6A
    h3.font.bold = False #Define se vai ser negrito ou não

    tabela = docum.add_table(rows = len(processos)+1, cols = 4, style='Table Grid')
    Format = tabela.rows[0].cells
    Format[0].text = 'Processos'
    Format[1].text = 'Tempo de Execução'
    Format[2].text= 'Ordem de Chegada'
    Format[3].text= 'Prioridade'
    tabela.aligment = WD_TABLE_ALIGNMENT.CENTER
    for x in range(1, len(processos)+1):
        Format = tabela.rows[x].cells
        Format[0].text = str(processos[x-1][0])
        Format[1].text = str(processos[x-1][1])
        Format[2].text = str(processos[x-1][2])
        Format[3].text = str(processos[x-1][3])
    docum.add_paragraph('', style="H2")
    docum.add_picture('image.png', width=Inches(4))
    val = 0
    for x in y:
        docum.add_paragraph('Tempo de Espera do Processo '+str(x)+': '+str(xinicio[val])+' NanosSegundos', style="H3")
        val+=1
    docum.add_paragraph('Tempo médio de espera: {0:.2f} ns'.format(valuecalcu[4]), style="H2")
    docum.add_paragraph('Tempo Total: {0:.2f} ns'.format(valuecalcu[2]), style="H2")
    docum.add_paragraph('Código do relatório: {0}'.format(codigo), style="H3")
    try:
        DiretorioSelecionado = os.getcwd()+'\SQLDate'
        dirfinaldocx = DiretorioSelecionado+'\RelatorioPIPESQL.docx'
        dirfinalpdf = DiretorioSelecionado+'\RelatorioPIPESQL.pdf'
        docum.save(dirfinaldocx)
        docx2pdf.convert(r"{0}".format(dirfinaldocx), r"{0}".format(dirfinalpdf))
        os.remove(dirfinaldocx)
        salvou = True
    except Exception:
        salvou = False
    os.remove('image.png')
    valorbytes = None
    try:
        arqpdf = open(dirfinalpdf, 'rb')
        valor =  arqpdf.readlines()
        arqpdf.close()
        valorbytes = b''.join(valor) #VARIÁVEL BYTE RESPONSÁVEL POR AMAZENAR DADOS NO SQL

        msq = QMessageBox()
        msq.setWindowTitle("Salvar Dados") #Define o titulo
        msq.setText('Você está prestes a salvar este dado no banco!') #Define descrição
        msq.setInformativeText("\nDeseja mesmo salvar?") #Define descrição secundária
        msq.setIcon(QMessageBox.Warning) #Abre mensagem de atenção
        msq.setStandardButtons(QMessageBox.Yes|QMessageBox.Cancel) #Criar opções de click na mensagem
        resposta = msq.exec_()
        if resposta == QMessageBox.Yes:
            try:
                con = conexaobanco('WILLIANQUIRINO\SQLEXPRESS', 'PIPESQL')
                cursor = con.cursor()
                cursor.execute("INSERT INTO BancoPipe(COD_Pipe, PDF_Pipe) VALUES(?, ?)", (codigo, valorbytes)) 
                cursor.commit()
                cursor.close()

                msq2 = QMessageBox()
                msq2.setWindowTitle("Dados Salvos") #Define o titulo
                msq2.setText('O seu dado está salvo com sucesso\nNúmero de protocolo: {0}\n\nSalve o seu número de protocolo para acessa-lo posteriormente!'.format(codigo)) #Define descrição
                #msq2.setInformativeText("\nDeseja mesmo salvar?") #Define descrição secundária
                msq2.setIcon(QMessageBox.Information) #Abre mensagem de atenção
                msq2.setTextInteractionFlags(Qt.TextSelectableByMouse) # PERMITE SELECIONAR O TEXTO NO MESSAGEBOX
                msq2.exec_()
                #QMessageBox.about(processo, "Dados Salvo com Sucesso", "Esse relatório foi salvo com sucesso\n\nO código do relatório é: {0}\nSalve o código para poder acessa-lo novamente!".format(codigo))
            except Exception as er:
                QMessageBox.about(processo, "Erro ao Salvar", "Não foi possível salvar este dado ao banco\n\nContate o desenvolvedor\nError: {0}".format(er))
   
        '''
        doguinho = open('FUNCIONACARALHO.pdf', 'wb')
        doguinho.write(valorbytes)
        doguinho.close()
        '''

    except Exception as err:
        salvou = False
        print(err)

    os.remove(dirfinalpdf)
    ps.close()
    return salvou

def criarcodigo() -> str:
    con = conexaobanco('WILLIANQUIRINO\SQLEXPRESS', 'PIPESQL')
    cursor = con.cursor()
    cursor.execute("SELECT max(ID_Pipe) from BancoPipe")
    row = cursor.fetchone()
    value: str
    for x in row:
        value = str(x)
    letras = 'abcdefghkijlmnopqrstuvwyxzABCDEFGHIJKLMNOPQRSTUVXWYZ123456789'
    for x in range(22):
        value += letras[random.randrange(0, len(letras))]
    return value
'''
processos = [
    ('P1', 13, 1, 17),
    ('P2', 54, 2, 26),
    ('P3', 24, 3, 229),
    ('P4', 35, 4, 114),
    ('P5', 14, 5, 62),
    ('P6', 23, 6, 68),
    ('P7', 3, 7, 105)           
]
valor = GerarRelatorioSQL(processos, devs.calcularFIFO(processos), devs.GerarTabela(devs.calcularFIFO(processos), False), 'FIFO', 'Nenhum')
'''

