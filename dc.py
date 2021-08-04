# PIP INSTALL PYTHON-DOCX
from typing import Counter
from docx import Document
import matplotlib.pyplot as ps
import devs
import numpy as np
###### BIBLIOTECA DE ESTILOS ###
from docx.enum.style import WD_STYLE_TYPE      # Define estilo para paragrafo
from docx.enum.table import WD_TABLE_ALIGNMENT # Define estilo para paragrafo
from docx.enum.text import WD_ALIGN_PARAGRAPH  #Define alinhamento de paragráfo
from docx.shared import Pt                     #Define tamanho da fonte
from docx.shared import RGBColor               #Define a cor da fonte
from docx.shared import Inches                 #Define tamanho

def resumo():
    doc = Document() #Cria uma variável do timo Document em doc

    doc.add_heading("Titulo", 3) #Adiciona titulo para o documento ("Nome", tamanho de formatação)
    doc.add_paragraph('Adicione uma linha aqui') #Adiciona um parágrafo <uma linha>

    ###### Modificar estilos das linhas e paragrafo ##########
    estilos = doc.styles         # Passa styles de doc para variável estilo
    paragrafo = estilos.add_style('Paragrafo', WD_ALIGN_PARAGRAPH)
    paragrafo.font.name = 'Calibri' #Define o tipo de fonte
    paragrafo.size = Pt(11) # Define tamanho da fonte
    paragrafo.base_style = estilos["Heading 2"] # Usa a base de style de um determinada variável
    paragrafo.font.color.rgb = RGBColor(79, 129, 189) # Define a cor da fonte
    paragrafo.font.bold = True #Define se a fonte deve ou não ser negrito

    ######## Acessando os CAMPOS e Formatando ################
    doc.add_paragraph('Olá', style = 'Paragrafo') # Adiciona o nome do estilo criado em estilos.add_style()
    doc.add_picture('image.png', width=Inches(1)) # Adiciona uma imagem (URL da IMAGEM) width = Inches(tamanho) "DEFINE LARGURA"

    ######## FORMATANDO TABELAS ##############
    tabela = doc.add_table(rows = 1, cols = 3, style='Table Grid') #Adiciona uma tabela (rows = 'Numeros de linhas', cols = 'Numeros de colunas', styles= 'Tipo de estilo da tabela')
    tabela.aligment = WD_TABLE_ALIGNMENT.CENTER # Define centralização, esquerdo ou direito do word
    tabela.autofit = False # Deixar tabela esticada ou não
    #rows - linha  /  cols - Colunas#
    FormatarTabela = tabela.rows[0].cells
    FormatarTabela[0].text = 'Coluna 1'
    # Style Paragraph
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)

    #Style Heading 2
    h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(79, 129, 189)
    h2.font.bold = False #Define se vai ser negrito ou não

    #Style Heading 3
    h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(79, 129, 189)
    h3.font.bold = False #Define se vai ser negrito ou não

    p = document.add_paragraph()
    r = p.add_run()
    r.add_text('Calculo de Escalonamento: Relatório')
    r.add_picture('PipesTransparent.png', width=Inches(1.25))

    doc.save('Teste.docx') #Salva o documento (Diretório)

processos = [('P1', 7, 1, 2),('P2', 12, 2, 3),('P3', 40, 3, 4),('P4', 25, 4, 2),('P5', 36, 5, 1),('P6', 20, 6, 3)]


valuecalcu = devs.calcularFIFO(processos)
valor = devs.GerarTabela(valuecalcu, False)
    # VARIAVEIS DE CONSTRUÇÃO DO GRÁFICO
y = valor[0] # PROCESSOS DA TABELA Y
c = []  # TIPO USADO (quantidade de processos)
for x in y:
    c.append(1)
xinicio = valor[1] #Inicio processo
xfim = valor[2] # Fim processo
color_mapper = np.vectorize(lambda x: {0: 'red', 1: 'blue'}.get(x))
ps.hlines(y, xinicio, xfim, colors=color_mapper(c), lw=25)
ps.ylabel('PROCESSOS')
ps.xlabel('Tempo de Execução (ns)')
ps.grid(True) #Deixar tabelado
ps.savefig('image.png')

docum = Document()

docum.add_heading('Relatório: Cálculo de Escalonamento', 0)
docum.add_heading('Tipo de escalonamento: FIFO\nCritério de Desempate: Nenhum')
styles = docum.styles
# Style Paragraph
p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
p.font.name = "Calibri"
p.font.size = Pt(11)
#Style Heading 2
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.font.name = "Calisto MT"
h2.font.size = Pt(12)
h2.font.bold = True #Define se vai ser negrito ou não
#Style Heading 3
h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.font.name = "Calisto MT"
h3.font.size = Pt(10)
h3.font.color.rgb = RGBColor(11, 0x0A, 0x6A) ##110A6A
h3.font.bold = False #Define se vai ser negrito ou não

tabela = docum.add_table(rows = len(processos)+1, cols = 4, style='Table Grid')
Format = tabela.rows[0].cells
Format[0].text = 'Processos'
Format[1].text = 'Tempo de Execução'
Format[2].text= 'Ordem de Chegada'
Format[3].text= 'Prioridade'
tabela.aligment = WD_TABLE_ALIGNMENT.CENTER
for x in range(1, len(processos)+1):
    Format = tabela.rows[x].cells
    Format[0].text = str(processos[x-1][0])
    Format[1].text = str(processos[x-1][1])
    Format[2].text = str(processos[x-1][2])
    Format[3].text = str(processos[x-1][3])
docum.add_picture('image.png', width=Inches(4))
val = 0
for x in y:
    docum.add_paragraph('Tempo de Espera do Processo '+str(x)+': '+str(xinicio[val])+' NanosSegundos', style="H3")
    val+=1
docum.add_paragraph('Tempo médio de espera: {0:.2f} ns'.format(valuecalcu[4]), style="H2")
docum.add_paragraph('Tempo Total: {0:.2f} ns'.format(valuecalcu[2]), style="H2")
docum.save('tabela.docx')